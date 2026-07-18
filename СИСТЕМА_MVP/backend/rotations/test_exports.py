from datetime import date, timedelta
from io import BytesIO

from django.test import TestCase
from django.utils import timezone
from docx import Document
from openpyxl import load_workbook

from shifts.models import ShiftType, WatchPeriod
from users.models import Employee

from .documents import PACKET_MARKER, build_extension_data_packet, document_bytes
from .exports import build_cycle_workbook, workbook_bytes
from .models import RotationCollectionCycle, RotationResponse, WatchExtensionCase


class RotationExportTests(TestCase):
    @classmethod
    def setUpTestData(cls):
        cls.generated_by = Employee.objects.create(full_name="Табельщик Тестовый")
        cls.route_employee = Employee.objects.create(full_name="Сотрудник с маршрутом")
        cls.approved_employee = Employee.objects.create(full_name="Одобренный Сотрудник")
        cls.rejected_employee = Employee.objects.create(full_name="Отклоненный Сотрудник")
        cls.site_manager = Employee.objects.create(full_name="Начальник Участка")

        cls.watch_period = WatchPeriod.objects.create(
            name="Вахта август–сентябрь",
            starts_on=date(2026, 8, 15),
            ends_on=date(2026, 9, 14),
        )
        cls.now = timezone.now()
        cls.cycle = RotationCollectionCycle.objects.create(
            name="Перевахта на август",
            target_watch_period=cls.watch_period,
            response_deadline=cls.now + timedelta(days=3),
            status=RotationCollectionCycle.Status.DRAFT,
            revision=2,
            created_by=cls.generated_by,
        )

        cls.route_response = RotationResponse.objects.create(
            cycle=cls.cycle,
            employee=cls.route_employee,
            snapshot_full_name='=HYPERLINK("https://example.invalid";"Открыть")',
            snapshot_personnel_number="Т-001",
            snapshot_position="Водитель автомобиля",
            snapshot_department="Горный участок",
            snapshot_work_schedule="Вахтовый график",
            snapshot_brigade_number=1,
            state=RotationResponse.State.SUBMITTED,
            intent=RotationResponse.Intent.ARRIVAL,
            next_shift_type=ShiftType.DAY,
            shift_source=RotationResponse.ShiftSource.EMPLOYEE,
            departure_on=date(2026, 8, 13),
            arrival_on=date(2026, 8, 15),
            route_text="+Хабаровск — Малмыж",
            travel_mode=RotationResponse.TravelMode.BUS,
            transfer_mode=RotationResponse.TransferMode.ORGANIZED,
            transport_details="Автобус от пункта сбора",
            comment="@проверить список",
            submitted_by=cls.route_employee,
            submitted_at=cls.now,
        )
        cls.approved_response = RotationResponse.objects.create(
            cycle=cls.cycle,
            employee=cls.approved_employee,
            snapshot_full_name="Одобренный Сотрудник",
            snapshot_personnel_number="Т-002",
            snapshot_position="Машинист экскаватора",
            snapshot_department="Горный участок",
            snapshot_work_schedule="Вахтовый график",
            snapshot_brigade_number=2,
            state=RotationResponse.State.SUBMITTED,
            intent=RotationResponse.Intent.EXTENSION,
            next_shift_type=ShiftType.NIGHT,
            shift_source=RotationResponse.ShiftSource.ACTIVE_ASSIGNMENT,
            comment="Готов остаться на следующую вахту",
            submitted_by=cls.approved_employee,
            submitted_at=cls.now,
        )
        cls.rejected_response = RotationResponse.objects.create(
            cycle=cls.cycle,
            employee=cls.rejected_employee,
            snapshot_full_name="Отклоненный Сотрудник",
            snapshot_personnel_number="Т-003",
            snapshot_position="Водитель автомобиля",
            snapshot_department="Горный участок",
            snapshot_work_schedule="Вахтовый график",
            snapshot_brigade_number=3,
            state=RotationResponse.State.SUBMITTED,
            intent=RotationResponse.Intent.EXTENSION,
            next_shift_type="",
            shift_source=RotationResponse.ShiftSource.UNKNOWN,
            submitted_by=cls.rejected_employee,
            submitted_at=cls.now,
        )
        cls.approved_case = WatchExtensionCase.objects.create(
            response=cls.approved_response,
            extension_start=date(2026, 8, 15),
            extension_end=date(2026, 9, 14),
            decision_status=WatchExtensionCase.DecisionStatus.APPROVED,
            decision_by=cls.site_manager,
            decision_at=cls.now,
            decision_comment="Кандидат согласован",
        )
        cls.rejected_case = WatchExtensionCase.objects.create(
            response=cls.rejected_response,
            extension_start=date(2026, 8, 15),
            extension_end=date(2026, 9, 14),
            decision_status=WatchExtensionCase.DecisionStatus.REJECTED,
            decision_by=cls.site_manager,
            decision_at=cls.now,
            decision_comment="Не требуется на следующую вахту",
        )

    @staticmethod
    def _column_by_header(sheet, header):
        return next(
            cell.column
            for cell in sheet[4]
            if cell.value == header
        )

    @staticmethod
    def _document_text(document):
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.extend(cell.text for cell in row.cells)
        return "\n".join(parts)

    def test_workbook_has_four_sheets_and_only_approved_extension(self):
        workbook = build_cycle_workbook(self.cycle, generated_by=self.generated_by)

        self.assertEqual(
            workbook.sheetnames,
            ["Сводка", "Маршруты", "Продление к оформлению", "Контроль ответов"],
        )
        approved_sheet = workbook["Продление к оформлению"]
        approved_values = [
            str(cell.value or "")
            for row in approved_sheet.iter_rows(min_row=5)
            for cell in row
        ]
        self.assertIn("Одобренный Сотрудник", approved_values)
        self.assertNotIn("Отклоненный Сотрудник", approved_values)
        self.assertEqual(approved_sheet.max_row, 5)

    def test_workbook_neutralizes_formulas_and_keeps_real_dates(self):
        workbook = build_cycle_workbook(self.cycle, generated_by=self.generated_by)
        route_sheet = workbook["Маршруты"]

        full_name_column = self._column_by_header(route_sheet, "ФИО")
        route_column = self._column_by_header(route_sheet, "Маршрут")
        departure_column = self._column_by_header(route_sheet, "Дата выезда")
        self.assertTrue(route_sheet.cell(5, full_name_column).value.startswith("'="))
        self.assertTrue(route_sheet.cell(5, route_column).value.startswith("'+"))
        self.assertEqual(route_sheet.cell(5, departure_column).value, date(2026, 8, 13))
        self.assertEqual(route_sheet.cell(5, departure_column).number_format, "dd.mm.yyyy")

    def test_workbook_bytes_can_be_opened_again(self):
        payload = workbook_bytes(
            build_cycle_workbook(self.cycle, generated_by=self.generated_by)
        )

        reopened = load_workbook(BytesIO(payload))
        self.assertEqual(
            reopened.sheetnames,
            ["Сводка", "Маршруты", "Продление к оформлению", "Контроль ответов"],
        )
        departure_column = self._column_by_header(reopened["Маршруты"], "Дата выезда")
        self.assertIsInstance(
            reopened["Маршруты"].cell(5, departure_column).value,
            date,
        )

    def test_document_packet_is_marked_and_contains_only_approved_cases(self):
        payload = document_bytes(
            build_extension_data_packet(self.cycle, generated_by=self.generated_by)
        )

        reopened = Document(BytesIO(payload))
        text = self._document_text(reopened)
        self.assertIn(PACKET_MARKER, text)
        self.assertIn("Одобренный Сотрудник", text)
        self.assertIn("Т-002", text)
        self.assertNotIn("Отклоненный Сотрудник", text)
        self.assertNotIn("Т-003", text)
