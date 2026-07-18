from __future__ import annotations

from datetime import date, datetime
from io import BytesIO
import re

from django.core.exceptions import ObjectDoesNotExist
from django.utils import timezone
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

from .models import RotationCollectionCycle, RotationResponse, WatchExtensionCase


PACKET_MARKER = "ПАКЕТ ИСХОДНЫХ ДАННЫХ — НЕ ПРИКАЗ"
_INVALID_XML_CHARACTERS = re.compile(
    "[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]"
)


def _clean_text(value, *, empty="—"):
    if value in (None, ""):
        return empty
    text = _INVALID_XML_CHARACTERS.sub("", str(value))
    return text or empty


def _display(instance, field_name, *, default="—"):
    if instance is None:
        return default
    raw_value = getattr(instance, field_name, None)
    if raw_value in (None, ""):
        return default
    display_method = getattr(instance, f"get_{field_name}_display", None)
    if callable(display_method):
        displayed = display_method()
        if displayed not in (None, ""):
            return _clean_text(displayed, empty=default)
    return _clean_text(raw_value, empty=default)


def _format_value(value):
    if value in (None, ""):
        return "—"
    if isinstance(value, datetime):
        if timezone.is_aware(value):
            value = timezone.localtime(value)
        return value.strftime("%d.%m.%Y %H:%M")
    if isinstance(value, date):
        return value.strftime("%d.%m.%Y")
    return _clean_text(value)


def _person_name(person):
    if person is None:
        return "—"
    return _clean_text(getattr(person, "full_name", "") or str(person))


def _shift_label(response: RotationResponse):
    if not response.next_shift_type:
        return "Не определена"
    return _display(response, "next_shift_type", default="Не определена")


def _response_snapshot(response: RotationResponse, field_name):
    snapshot_name = {
        "full_name": "snapshot_full_name",
        "personnel_number": "snapshot_personnel_number",
        "position": "snapshot_position",
        "department": "snapshot_department",
    }[field_name]
    value = getattr(response, snapshot_name, None)
    if value in (None, ""):
        legacy_name = {
            "full_name": "employee_name",
            "personnel_number": "personnel_number",
            "position": "position_name",
            "department": "department_name",
        }[field_name]
        value = getattr(response, legacy_name, "")
    return value


def _extension_case(response: RotationResponse):
    try:
        return response.extension_case
    except ObjectDoesNotExist:
        return None


def _approved_pairs(cycle: RotationCollectionCycle):
    queryset = cycle.responses.select_related(
        "extension_case__decision_by",
        "extension_case__documentation_by",
    ).order_by("snapshot_full_name", "id")
    pairs = []
    for response in queryset:
        case = _extension_case(response)
        if case and case.decision_status == "approved":
            pairs.append((response, case))
    return pairs


def _set_document_defaults(document: Document):
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10)
    for section in document.sections:
        header = section.header
        paragraph = header.paragraphs[0]
        paragraph.text = PACKET_MARKER
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.runs[0]
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(174, 53, 46)


def _add_packet_title(document: Document, cycle, generated_by, approved_count):
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(PACKET_MARKER)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(174, 53, 46)

    warning = document.add_paragraph()
    warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
    warning_run = warning.add_run(
        "Документ содержит только проверяемые исходные данные. "
        "Он не является письменным согласием, приказом или иным кадровым документом."
    )
    warning_run.bold = True
    warning_run.font.name = "Arial"
    warning_run.font.size = Pt(10)

    target_watch = cycle.target_watch_period
    metadata = document.add_table(rows=0, cols=2)
    metadata.style = "Table Grid"
    metadata_rows = (
        ("Сбор", cycle.name),
        ("Целевая вахта", target_watch.name),
        ("Период вахты", f"{_format_value(target_watch.starts_on)} — {_format_value(target_watch.ends_on)}"),
        ("Ревизия", cycle.revision),
        ("Сформировал", _person_name(generated_by)),
        ("Сформировано", _format_value(timezone.now())),
        ("Одобренных заявок", approved_count),
    )
    for label, value in metadata_rows:
        cells = metadata.add_row().cells
        cells[0].text = _clean_text(label)
        cells[1].text = _format_value(value)
        cells[0].paragraphs[0].runs[0].bold = True


def _add_case_table(document: Document, number, response, case):
    heading = document.add_heading(
        f"{number}. {_clean_text(_response_snapshot(response, 'full_name'))}",
        level=1,
    )
    heading.runs[0].font.name = "Arial"
    heading.runs[0].font.size = Pt(13)

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    rows = [
        ("ФИО", _response_snapshot(response, "full_name")),
        ("Должность", _response_snapshot(response, "position")),
        ("Подразделение", _response_snapshot(response, "department")),
    ]
    personnel_number = _response_snapshot(response, "personnel_number")
    if personnel_number:
        rows.append(("Табельный номер", personnel_number))
    rows.extend(
        [
            ("Смена", _shift_label(response)),
            ("Продление с", case.extension_start),
            ("Продление по", case.extension_end),
            ("Решение", _display(case, "decision_status")),
            ("Решение принял", _person_name(case.decision_by)),
            ("Дата решения", case.decision_at),
            ("Комментарий решения", case.decision_comment),
            ("Статус подготовки документов", _display(case, "documentation_status")),
            ("Примечание к документам", case.documentation_note),
        ]
    )
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = _clean_text(label)
        cells[1].text = _format_value(value)
        cells[0].paragraphs[0].runs[0].bold = True

    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = note.add_run(PACKET_MARKER)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(174, 53, 46)


def build_extension_data_packet(
    cycle: RotationCollectionCycle,
    generated_by=None,
):
    approved_pairs = _approved_pairs(cycle)
    document = Document()
    _set_document_defaults(document)
    _add_packet_title(document, cycle, generated_by, len(approved_pairs))

    if not approved_pairs:
        paragraph = document.add_paragraph(
            "В выбранном сборе нет одобренных заявок на продление."
        )
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return document

    for number, (response, case) in enumerate(approved_pairs, start=1):
        document.add_page_break()
        _add_case_table(document, number, response, case)
    return document


def document_bytes(document: Document):
    output = BytesIO()
    document.save(output)
    return output.getvalue()
